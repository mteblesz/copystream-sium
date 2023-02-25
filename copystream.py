import pyperclip as pp
from docx import Document 
import signal
import sys
from datetime import datetime


FONT_NAME = 'Arial'


start_text = \
'''
copystream
author: Michał Szeląg (MIT Licence)

'''
mainloop_text = \
'''-----------------------------------------------------------
Kopiowany tekst będzie zapisywany do pliku.
Kliknij Ctrl+c w tym oknie terminala, 
aby zapauzować pracę programu i wywołać menu.
'''
menu_text = \
'''-----------------------------------------------------------
Program zapauzowano.
Wpisz odpowiednią literę opcji i zatwierdź klawiszem Enter:
    q - wyjdź z programu
    c - kontynuuj pracę (wyjście z menu)
    n - utwórz nowy plik .docx
    h - wyświetl pomoc
'''
help_text = \
'''
Pomoc:
Po uruchomieniu programu każdy skopiowany przez użytkownika tekst będzie zapisywany do pliku. 
Plik, w formacie .docx, otrzymuje nazwę będącą zapisaną datą i godziną uruchomienia programu.
Użytkownik może skorzystać z menu pauzy za pomocą skrótu klawiszowego Ctrl+c, kiedy wybrane jest okno z uruchomionym programem. W przeciwnym wypadku po prostu skopiuje nowy tekst.
W menu można wyjść z programu (postęp jest zapisywany na bieżąco, nie należy się obawiać jego utraty), wrócić do trybu zbierania tekstu lub poprzez utworzenie nowego pliku wyjściowego wygodnie rozdzielić pracę.
Kontakt do autora: https://github.com/teblesz
Program udostępniony jest na licencji MIT.
'''

run = True
document = Document()
docx_name = ''

def get_new_doc_name():
    return datetime.now().strftime("%Y.%m.%d_%H-%M-%S") + '.docx'

def quit():
    global run
    run = False
    print("Program kończy pracę...")
    sys.exit(0)

def make_new_doc():
    global document
    global docx_name
    document = Document()
    docx_name = get_new_doc_name()
    document.save(docx_name)

def sigint_handler(signum, frame):
    signal.signal(signal.SIGINT, signal.SIG_IGN)
    print(menu_text)
    while True:
        try:
            opt = input()
            if opt == 'q':
                quit()
            elif opt == 'c':
                break
            elif opt == 'n':
                make_new_doc()
            elif opt == 'h':
                print(help_text)
            else:
                print("błędna opcja")
        except EOFError:
            print("błędna opcja")
    print('Wznawianie pracy...')
    print(mainloop_text)
    signal.signal(signal.SIGINT, sigint_handler)


print(start_text)
signal.signal(signal.SIGINT, sigint_handler)
make_new_doc()
print(mainloop_text)
while run:
    text = pp.waitForNewPaste()
    print(text)
    document = Document(docx_name) # reset buffer
    par = document.add_paragraph(text)
    document.styles['Normal'].font.name = FONT_NAME
    par.style = document.styles['Normal']
    document.save(docx_name)